from pathlib import Path
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# =========================================================
# PROJECT PATHS
# =========================================================

BASE_DIR = Path(__file__).resolve().parent.parent

DATA_DIR = BASE_DIR / "data"
TABLES_DIR = BASE_DIR / "results" / "tables"
FIGURES_DIR = BASE_DIR / "results" / "figures"
REPORTS_DIR = BASE_DIR / "reports"

REPORTS_DIR.mkdir(parents=True, exist_ok=True)

DATA_PATH = DATA_DIR / "final_malnutrition_dataset.csv"

OUTPUT_PATH = REPORTS_DIR / "Malnutrition_Research_Automated_Report.docx"


# =========================================================
# START
# =========================================================

print("\n======================================")
print("     AUTOMATED RESEARCH REPORT")
print("======================================")


# =========================================================
# LOAD DATA
# =========================================================

print("\nLoading final dataset...")

df = pd.read_csv(DATA_PATH)

print("Dataset loaded successfully.")
print(f"Rows: {df.shape[0]}")
print(f"Columns: {df.shape[1]}")


# =========================================================
# CREATE WORD DOCUMENT
# =========================================================

document = Document()


# =========================================================
# PAGE SETTINGS
# =========================================================

section = document.sections[0]

section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)


# =========================================================
# FONT
# =========================================================

document.styles["Normal"].font.name = "Times New Roman"
document.styles["Normal"].font.size = Pt(11)


# =========================================================
# HELPER FUNCTION FOR CSV TABLE
# =========================================================

def add_csv_table(filename, heading, max_rows=100):

    file_path = TABLES_DIR / filename

    if not file_path.exists():

        document.add_paragraph(
            f"{filename} was not found."
        )

        return

    try:

        data = pd.read_csv(file_path)

        document.add_heading(
            heading,
            level=2
        )

        if data.empty:

            document.add_paragraph(
                "No records available."
            )

            return

        data = data.head(max_rows)

        table = document.add_table(
            rows=1,
            cols=len(data.columns)
        )

        table.style = "Table Grid"

        table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )

        # Header

        for i, column in enumerate(data.columns):

            table.rows[0].cells[i].text = str(
                column
            )

        # Data rows

        for _, row in data.iterrows():

            cells = table.add_row().cells

            for i, value in enumerate(row):

                if pd.isna(value):

                    text = ""

                else:

                    text = str(value)

                cells[i].text = text

    except Exception as e:

        document.add_paragraph(
            f"Unable to load {filename}: {e}"
        )


# =========================================================
# TITLE
# =========================================================

title = document.add_paragraph()

title.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = title.add_run(
    "Automated Malnutrition Research Report"
)

run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(18)


subtitle = document.add_paragraph()

subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = subtitle.add_run(
    "Statistical Analysis and Model Evaluation"
)

run.italic = True
run.font.size = Pt(12)


document.add_paragraph("")


# =========================================================
# 1. INTRODUCTION
# =========================================================

document.add_heading(
    "1. Introduction",
    level=1
)

document.add_paragraph(
    "This report presents an automated statistical analysis "
    "of child malnutrition using a practice dataset. The "
    "analysis was conducted through a reproducible Python "
    "research automation pipeline."
)

document.add_paragraph(
    "The workflow includes data import, validation, cleaning, "
    "missing value assessment, descriptive statistics, "
    "visualization, correlation analysis, multicollinearity "
    "assessment, logistic regression, model diagnostics, "
    "publication tables, publication figures, and automated "
    "report generation."
)


# =========================================================
# 2. DATASET INFORMATION
# =========================================================

document.add_heading(
    "2. Dataset Information",
    level=1
)

document.add_paragraph(
    f"The final analytical dataset contains "
    f"{df.shape[0]} observations and "
    f"{df.shape[1]} variables."
)

document.add_paragraph(
    "The three outcome variables are stunting, wasting, "
    "and underweight."
)

document.add_paragraph(
    "Numerical predictor variables include age in months, "
    "birth weight, dietary diversity, and maternal BMI."
)

document.add_paragraph(
    "Categorical predictor variables include sex, maternal "
    "education, wealth index, urban or rural residence, "
    "sanitation, water source, and immunization."
)


# =========================================================
# 3. DATA QUALITY
# =========================================================

document.add_heading(
    "3. Data Quality Assessment",
    level=1
)

missing_total = int(
    df.isna().sum().sum()
)

duplicate_total = int(
    df.duplicated().sum()
)

document.add_paragraph(
    f"Total missing values: {missing_total}."
)

document.add_paragraph(
    f"Duplicate records: {duplicate_total}."
)

if missing_total == 0:

    document.add_paragraph(
        "No missing values were detected in the final "
        "dataset. Therefore, no imputation was required."
    )

if duplicate_total == 0:

    document.add_paragraph(
        "No duplicate records were detected."
    )


# =========================================================
# 4. DESCRIPTIVE STATISTICS
# =========================================================

document.add_heading(
    "4. Descriptive Statistics",
    level=1
)

document.add_paragraph(
    "Descriptive statistics were calculated for numerical "
    "variables, while frequency and percentage distributions "
    "were calculated for categorical variables."
)

add_csv_table(
    "Table_1_Numeric_Descriptive_Statistics.csv",
    "Table 1. Numeric Descriptive Statistics"
)

add_csv_table(
    "Table_1B_Categorical_Descriptive_Statistics.csv",
    "Table 1B. Categorical Descriptive Statistics",
    max_rows=100
)


# =========================================================
# 5. CORRELATION ANALYSIS
# =========================================================

document.add_heading(
    "5. Correlation Analysis",
    level=1
)

document.add_paragraph(
    "Pearson correlation coefficients were calculated among "
    "the numerical predictor variables."
)

add_csv_table(
    "Table_2_Correlation_Matrix.csv",
    "Table 2. Correlation Matrix"
)

strong_file = TABLES_DIR / "strong_correlations.csv"

if strong_file.exists():

    strong_df = pd.read_csv(
        strong_file
    )

    if strong_df.empty:

        document.add_paragraph(
            "No pair of numerical variables showed an "
            "absolute correlation of 0.70 or greater."
        )


# =========================================================
# 6. MULTICOLLINEARITY
# =========================================================

document.add_heading(
    "6. Multicollinearity Assessment",
    level=1
)

document.add_paragraph(
    "Variance Inflation Factor (VIF) was used to assess "
    "multicollinearity among numerical predictor variables. "
    "A VIF value greater than 5 was considered indicative "
    "of potential multicollinearity."
)

add_csv_table(
    "Table_3_VIF_Analysis.csv",
    "Table 3. VIF Analysis"
)


# =========================================================
# 7. LOGISTIC REGRESSION
# =========================================================

document.add_heading(
    "7. Logistic Regression Analysis",
    level=1
)

document.add_paragraph(
    "Separate logistic regression models were fitted for "
    "stunting, wasting, and underweight."
)

document.add_paragraph(
    "The models included numerical and categorical predictor "
    "variables defined in the research automation pipeline."
)

add_csv_table(
    "Table_4_Logistic_Regression.csv",
    "Table 4. Logistic Regression Results",
    max_rows=200
)


# =========================================================
# 8. MODEL DIAGNOSTICS
# =========================================================

document.add_heading(
    "8. Model Diagnostics",
    level=1
)

document.add_paragraph(
    "Model performance was evaluated using accuracy, "
    "sensitivity, specificity, ROC-AUC, and confusion "
    "matrix measures."
)

add_csv_table(
    "Table_5_Model_Diagnostics.csv",
    "Table 5. Model Diagnostic Results"
)


# =========================================================
# 9. ROC-AUC
# =========================================================

document.add_heading(
    "9. ROC-AUC Analysis",
    level=1
)

roc_file = TABLES_DIR / "roc_auc_summary.csv"

if roc_file.exists():

    add_csv_table(
        "roc_auc_summary.csv",
        "Table 6. ROC-AUC Summary"
    )


# =========================================================
# 10. KEY NUMERICAL FINDINGS
# =========================================================

document.add_heading(
    "10. Key Findings",
    level=1
)


# ---------------------------------------------------------
# Outcome prevalence
# ---------------------------------------------------------

if "stunting" in df.columns:

    stunting_numeric = (
        df["stunting"]
        .astype(str)
        .str.strip()
        .str.lower()
        .eq("yes")
        .astype(float)
    )

    stunting_yes = stunting_numeric.mean() * 100

    document.add_paragraph(
        f"Stunting prevalence was {stunting_yes:.2f}%."
    )


if "wasting" in df.columns:

    wasting_numeric = (
        df["wasting"]
        .astype(str)
        .str.strip()
        .str.lower()
        .eq("yes")
        .astype(float)
    )

    wasting_yes = wasting_numeric.mean() * 100

    document.add_paragraph(
        f"Wasting prevalence was {wasting_yes:.2f}%."
    )


if "underweight" in df.columns:

    underweight_numeric = (
        df["underweight"]
        .astype(str)
        .str.strip()
        .str.lower()
        .eq("yes")
        .astype(float)
    )

    underweight_yes = underweight_numeric.mean() * 100

    document.add_paragraph(
        f"Underweight prevalence was "
        f"{underweight_yes:.2f}%."
    )


# ---------------------------------------------------------
# Numerical summaries
# IMPORTANT:
# Explicit numeric conversion prevents string mean error
# ---------------------------------------------------------

if "age_months" in df.columns:

    age_values = pd.to_numeric(
        df["age_months"],
        errors="coerce"
    )

    age_mean = age_values.mean()

    document.add_paragraph(
        f"Mean child age was {age_mean:.2f} months."
    )


if "birth_weight" in df.columns:

    birth_weight_values = pd.to_numeric(
        df["birth_weight"],
        errors="coerce"
    )

    birth_weight_mean = birth_weight_values.mean()

    document.add_paragraph(
        f"Mean birth weight was "
        f"{birth_weight_mean:.2f} kg."
    )


if "dietary_diversity" in df.columns:

    dietary_values = pd.to_numeric(
        df["dietary_diversity"],
        errors="coerce"
    )

    dietary_mean = dietary_values.mean()

    document.add_paragraph(
        f"Mean dietary diversity score was "
        f"{dietary_mean:.2f}."
    )


if "maternal_bmi" in df.columns:

    bmi_values = pd.to_numeric(
        df["maternal_bmi"],
        errors="coerce"
    )

    bmi_mean = bmi_values.mean()

    document.add_paragraph(
        f"Mean maternal BMI was "
        f"{bmi_mean:.2f}."
    )


# =========================================================
# 11. MODEL PERFORMANCE
# =========================================================

document.add_heading(
    "11. Model Performance",
    level=1
)

diagnostic_file = TABLES_DIR / "model_diagnostics.csv"

if diagnostic_file.exists():

    diagnostics = pd.read_csv(
        diagnostic_file
    )

    if not diagnostics.empty:

        # Convert ONLY required columns to numeric

        numeric_columns = [
            "Accuracy",
            "Sensitivity",
            "Specificity",
            "ROC_AUC"
        ]

        for col in numeric_columns:

            if col in diagnostics.columns:

                diagnostics[col] = pd.to_numeric(
                    diagnostics[col],
                    errors="coerce"
                )

        if (
            "ROC_AUC" in diagnostics.columns
            and diagnostics["ROC_AUC"].notna().any()
        ):

            best_index = diagnostics[
                "ROC_AUC"
            ].idxmax()

            best_row = diagnostics.loc[
                best_index
            ]

            outcome = str(
                best_row.get(
                    "Outcome",
                    "model"
                )
            )

            auc = float(
                best_row["ROC_AUC"]
            )

            document.add_paragraph(
                f"The highest ROC-AUC among the evaluated "
                f"models was observed for {outcome}, "
                f"with an ROC-AUC of {auc:.4f}."
            )


# =========================================================
# 12. FIGURES
# =========================================================

document.add_heading(
    "12. Figures",
    level=1
)

figures = [

    (
        "Figure_1_Outcome_Prevalence.png",
        "Figure 1. Outcome Prevalence"
    ),

    (
        "Figure_2_Age_Distribution.png",
        "Figure 2. Age Distribution"
    ),

    (
        "Figure_3_Birth_Weight_Distribution.png",
        "Figure 3. Birth Weight Distribution"
    ),

    (
        "Figure_4_Dietary_Diversity.png",
        "Figure 4. Dietary Diversity Distribution"
    ),

    (
        "Figure_5_Maternal_BMI_Distribution.png",
        "Figure 5. Maternal BMI Distribution"
    ),

    (
        "Figure_ROC_stunting.png",
        "Figure 6. ROC Curve for Stunting"
    ),

    (
        "Figure_ROC_wasting.png",
        "Figure 7. ROC Curve for Wasting"
    ),

    (
        "Figure_ROC_underweight.png",
        "Figure 8. ROC Curve for Underweight"
    ),

    (
        "Figure_9_ROC_AUC_Comparison.png",
        "Figure 9. ROC-AUC Comparison"
    )
]


for filename, caption in figures:

    figure_path = FIGURES_DIR / filename

    if figure_path.exists():

        paragraph = document.add_paragraph()

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(
            caption
        )

        run.bold = True

        image_paragraph = document.add_paragraph()

        image_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        image_paragraph.add_run().add_picture(
            str(figure_path),
            width=Inches(5.8)
        )


# =========================================================
# 13. CONCLUSION
# =========================================================

document.add_heading(
    "13. Conclusion",
    level=1
)

document.add_paragraph(
    "The research automation pipeline successfully completed "
    "the major stages of data preparation, exploratory "
    "analysis, statistical modelling, model diagnostics, "
    "and research output generation."
)

document.add_paragraph(
    "The automated workflow produces reproducible tables, "
    "figures, diagnostic results, and a structured research "
    "report from the processed dataset."
)

document.add_paragraph(
    "Because this is a practice dataset, the findings should "
    "not be interpreted as estimates of the actual prevalence "
    "or determinants of malnutrition in a real population."
)


# =========================================================
# 14. REPRODUCIBILITY
# =========================================================

document.add_heading(
    "14. Reproducibility Note",
    level=1
)

document.add_paragraph(
    "All analyses were implemented using Python scripts "
    "within the Research Automation project structure. "
    "The workflow is designed so that the analysis can be "
    "repeated after replacing the input dataset with an "
    "appropriately structured dataset."
)


# =========================================================
# SAVE DOCUMENT
# =========================================================

print("\nSaving Word report...")

document.save(
    OUTPUT_PATH
)


# =========================================================
# SUCCESS
# =========================================================

print("\n======================================")
print("   AUTOMATED REPORT COMPLETED")
print("======================================")

print("\nReport saved to:")

print(OUTPUT_PATH)

print("\n======================================")
print("          SUCCESS")
print("======================================")